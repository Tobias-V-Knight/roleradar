# agents/resume_builder.py
# -----------------------------------------------------------------------------
# Two responsibilities:
#   1. generate_content() — ask GPT-4o to produce tailored bullets/projects/skills
#      as structured JSON, sourced strictly from the master CV.
#   2. build_docx() — open the base resume DOCX as a template and replace only
#      the dynamic sections (bullets, projects, skills) in-place, then return
#      the modified file as bytes for download.
# -----------------------------------------------------------------------------

import json

from foundry.client import OPENAI_STUB_MODE, get_azure_openai_config

SYSTEM_MSG = (
    "You are an expert resume writer. Given a master CV, a job description, and "
    "the candidate's base resume, rewrite only the experience bullet points, select "
    "the best projects, and curate the skills list to match the job description. "
    "Use ONLY content that appears in the master CV — never fabricate experience. "
    "Return ONLY valid JSON."
)

USER_TEMPLATE = """JOB DESCRIPTION:
{jd_text}

MASTER CV (all available experience — the only content you may draw from):
{cv_text}

BASE RESUME (for structure reference — company names, titles, and dates are fixed):
{base_resume_text}

Return JSON with exactly these fields:
{{
  "experience_bullets": {{
    "<SHORT company name matching the base resume>": ["<bullet>", "<bullet>", "<bullet>"]
  }},
  "projects": [
    {{"name": "<project name>", "description": "<1-2 sentence description from CV>"}}
  ],
  "skills_tools": "<comma-separated tools ordered by relevance to JD>",
  "skills_methods": "<comma-separated methods ordered by relevance to JD>"
}}

Rules:
- Use short company name keys (e.g. "CARLSON ANALYTICS LAB") that are substrings of
  the full lines in the base resume — this is how they get matched to the right section.
- Keep the same number of bullets as the base resume for each company.
- Pick 2-3 projects from the CV that best match the JD; do not invent new ones.
- Skills must only include items present in the master CV.
- Each bullet must start with a strong past-tense action verb and quantify impact.
"""


def generate_content(jd_text: str, cv_text: str, base_resume_text: str) -> dict:
    """Call GPT-4o to produce tailored resume content. Stubs gracefully if no creds."""
    if OPENAI_STUB_MODE:
        print("[STUB] resume_builder: returning mock content (no Azure OpenAI creds)")
        return _stub_content()

    cfg = get_azure_openai_config()
    try:
        from openai import AzureOpenAI

        client = AzureOpenAI(
            api_key=cfg["api_key"],
            api_version=cfg["api_version"],
            azure_endpoint=cfg["base_url"],
        )
        resp = client.chat.completions.create(
            model=cfg["model"],
            messages=[
                {"role": "system", "content": SYSTEM_MSG},
                {"role": "user", "content": USER_TEMPLATE.format(
                    jd_text=jd_text[:5000],
                    cv_text=cv_text[:8000],
                    base_resume_text=base_resume_text[:4000],
                )},
            ],
            temperature=0.3,
            response_format={"type": "json_object"},
        )
        data = json.loads(resp.choices[0].message.content)
        data["_stub"] = False
        return data
    except Exception as e:
        print(f"[WARN] resume_builder live call failed ({e}); returning stub")
        s = _stub_content()
        s["_error"] = str(e)
        return s


def _stub_content() -> dict:
    return {
        "experience_bullets": {
            "CARLSON ANALYTICS LAB": [
                "[STUB] Add Azure OpenAI credentials for real tailored bullets.",
                "[STUB] Built analytics pipelines to surface performance insights.",
                "[STUB] Evaluated causal impact of key initiatives via regression.",
            ]
        },
        "projects": [
            {
                "name": "[STUB] Sample Project",
                "description": "Add Azure OpenAI credentials for real GPT-4o content.",
            }
        ],
        "skills_tools": "[STUB] Python, SQL, Tableau",
        "skills_methods": "[STUB] Analytics, A/B Testing, Causal Inference",
        "_stub": True,
    }


# -----------------------------------------------------------------------------
# DOCX builder — in-place template modification
# -----------------------------------------------------------------------------

def _set_para_text(para, new_text: str):
    """Replace a paragraph's visible text while preserving its run formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _find_company_key(para_text: str, experience_bullets: dict):
    """Return the matching key from experience_bullets if it's a substring of para_text."""
    upper = para_text.upper()
    for key in experience_bullets:
        if key.upper() in upper:
            return key
    return None


def _all_paragraphs_in_order(doc):
    """
    Yield every paragraph in document body order, including those inside tables.
    python-docx's doc.paragraphs skips table cells, which many resume templates
    use to align company names (left) against dates (right).
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    def _from_element(el):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            yield Paragraph(el, doc)
        elif tag == "tbl":
            tbl = Table(el, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
        # Recurse into body-level containers (e.g. txbx)
        for child in el:
            child_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if child_tag not in ("p", "tbl"):
                yield from _from_element(child)

    for child in doc.element.body:
        yield from _from_element(child)


def build_docx(base_docx_bytes: bytes, content_json: dict) -> bytes:
    """
    Open the base resume DOCX as a template, replace the dynamic sections
    (experience bullets, projects, skills) in-place, return modified bytes.
    """
    from io import BytesIO
    from docx import Document

    experience_bullets: dict = content_json.get("experience_bullets", {})
    projects: list = content_json.get("projects", [])
    skills_tools: str = content_json.get("skills_tools", "")
    skills_methods: str = content_json.get("skills_methods", "")

    doc = Document(BytesIO(base_docx_bytes))

    # --- First pass: bucket paragraphs by section (table-aware) ---
    section = None
    exp_company_buckets: dict = {}   # company_key -> [bullet paras]
    current_company_key = None
    project_paras: list = []
    skill_paras: list = []

    SECTION_HEADERS = {"EDUCATION", "EXPERIENCE", "DATA SCIENCE PROJECTS", "TECHNICAL SKILLS"}

    for para in _all_paragraphs_in_order(doc):
        text = para.text.strip()
        # Normalise: strip leading backtick/apostrophe artefacts some DOCX exports add
        text = text.lstrip("`'‘’“”").strip()
        upper = text.upper()

        if upper in SECTION_HEADERS:
            section = upper
            current_company_key = None
            continue

        if not text:
            continue

        if section == "EXPERIENCE":
            if "●" in text:
                if current_company_key:
                    exp_company_buckets.setdefault(current_company_key, []).append(para)
            else:
                matched = _find_company_key(text, experience_bullets)
                if matched:
                    current_company_key = matched

        elif section == "DATA SCIENCE PROJECTS":
            project_paras.append(para)

        elif section == "TECHNICAL SKILLS":
            skill_paras.append(para)

    # --- Replace experience bullets ---
    for key, bullet_paras in exp_company_buckets.items():
        new_bullets = experience_bullets.get(key, [])
        for i, para in enumerate(bullet_paras):
            if i < len(new_bullets):
                _set_para_text(para, "● " + new_bullets[i])
            else:
                _set_para_text(para, "")

    # --- Replace project paragraphs ---
    for i, para in enumerate(project_paras):
        if i < len(projects):
            proj = projects[i]
            if len(para.runs) >= 2:
                para.runs[0].text = proj["name"] + ": "
                para.runs[1].text = proj["description"]
                for run in para.runs[2:]:
                    run.text = ""
            else:
                _set_para_text(para, proj["name"] + ": " + proj["description"])
        else:
            _set_para_text(para, "")

    # --- Replace skills lines ---
    for para in skill_paras:
        text = para.text.strip().lower()
        if text.startswith("tools:"):
            if len(para.runs) >= 2:
                para.runs[0].text = "Tools: "
                para.runs[1].text = skills_tools
                for run in para.runs[2:]:
                    run.text = ""
            else:
                _set_para_text(para, "Tools: " + skills_tools)
        elif text.startswith("methods:"):
            if len(para.runs) >= 2:
                para.runs[0].text = "Methods: "
                para.runs[1].text = skills_methods
                for run in para.runs[2:]:
                    run.text = ""
            else:
                _set_para_text(para, "Methods: " + skills_methods)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
