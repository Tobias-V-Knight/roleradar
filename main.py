# main.py
# -----------------------------------------------------------------------------
# FastAPI application: REST API + static frontend + APScheduler background job.
#
# Run with:  python main.py    (or: uvicorn main:app --reload)
# Then open: http://localhost:8000
# -----------------------------------------------------------------------------

import threading
from typing import Optional

from apscheduler.schedulers.background import BackgroundScheduler
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.responses import FileResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

import config
import database
from foundry.orchestrator import get_orchestrator

app = FastAPI(title="RoleRadar", version="1.0")

# A simple in-memory status object so the UI can poll scrape progress.
SCRAPE_STATUS = {"running": False, "last_run": None, "message": "idle"}

# A lock so two scrape triggers can't run the pipeline concurrently.
_scrape_lock = threading.Lock()


# -----------------------------------------------------------------------------
# Pydantic request bodies
# -----------------------------------------------------------------------------
class CompanyIn(BaseModel):
    name: str
    careers_url: str
    tier: int = 1


class RoleIn(BaseModel):
    keyword: str


class LocationIn(BaseModel):
    city: str
    state: str = "US"


class ResumeIn(BaseModel):
    resume_text: str


# -----------------------------------------------------------------------------
# Startup: init DB, build orchestrator, start scheduler
# -----------------------------------------------------------------------------
scheduler = BackgroundScheduler()


@app.on_event("startup")
def startup():
    database.init_db()
    # Build the orchestrator once (prints the stub/live banners).
    get_orchestrator()
    # Schedule a recurring full scrape every N hours.
    scheduler.add_job(
        _run_full_scrape_job,
        "interval",
        hours=config.SCRAPE_INTERVAL_HOURS,
        id="full_scrape",
        replace_existing=True,
    )
    scheduler.start()
    print(f"[RoleRadar] started. Scheduled full scrape every "
          f"{config.SCRAPE_INTERVAL_HOURS}h. Open http://localhost:8000")


@app.on_event("shutdown")
def shutdown():
    scheduler.shutdown(wait=False)


# -----------------------------------------------------------------------------
# Scrape runners (threaded so the HTTP request returns immediately)
# -----------------------------------------------------------------------------
def _run_full_scrape_job(only_company_id=None, note="scheduled"):
    """Run a scrape, guarding against overlapping runs and updating status."""
    if not _scrape_lock.acquire(blocking=False):
        print("[RoleRadar] scrape already running; skipping new trigger")
        return
    SCRAPE_STATUS.update(running=True, message=f"scraping ({note})...")
    try:
        result = get_orchestrator().run_full_scrape(
            only_company_id=only_company_id, note=note
        )
        SCRAPE_STATUS.update(
            running=False,
            last_run=database.latest_run(),
            message=f"done: {result['matches']} matches across "
                    f"{result['companies_scraped']} companies",
        )
    except Exception as e:
        SCRAPE_STATUS.update(running=False, message=f"error: {e}")
        print(f"[ERROR] scrape job failed: {e}")
    finally:
        _scrape_lock.release()


def _trigger_async(only_company_id=None, note="manual"):
    """Kick off a scrape in a background thread and return immediately."""
    t = threading.Thread(
        target=_run_full_scrape_job,
        kwargs={"only_company_id": only_company_id, "note": note},
        daemon=True,
    )
    t.start()


# -----------------------------------------------------------------------------
# Company endpoints
# -----------------------------------------------------------------------------
@app.get("/api/companies")
def get_companies():
    return database.list_companies()


@app.post("/api/companies")
def post_company(body: CompanyIn):
    new_id = database.add_company(body.name, body.careers_url, body.tier)
    return {"id": new_id}


@app.delete("/api/companies/{company_id}")
def remove_company(company_id: int):
    database.delete_company(company_id)
    return {"ok": True}


@app.patch("/api/companies/{company_id}/toggle")
def toggle_company(company_id: int):
    database.toggle_company(company_id)
    return {"ok": True}


# -----------------------------------------------------------------------------
# Role endpoints
# -----------------------------------------------------------------------------
@app.get("/api/roles")
def get_roles():
    return database.list_roles()


@app.post("/api/roles")
def post_role(body: RoleIn):
    database.add_role(body.keyword)
    return {"ok": True}


@app.delete("/api/roles/{role_id}")
def remove_role(role_id: int):
    database.delete_role(role_id)
    return {"ok": True}


# -----------------------------------------------------------------------------
# Location endpoints
# -----------------------------------------------------------------------------
@app.get("/api/locations")
def get_locations():
    return database.list_locations()


@app.post("/api/locations")
def post_location(body: LocationIn):
    database.add_location(body.city, body.state)
    return {"ok": True}


@app.delete("/api/locations/{location_id}")
def remove_location(location_id: int):
    database.delete_location(location_id)
    return {"ok": True}


# -----------------------------------------------------------------------------
# Scraping endpoints
# -----------------------------------------------------------------------------
@app.post("/api/scrape")
def scrape_all():
    """Trigger a full scrape of all active companies (async)."""
    _trigger_async(note="full")
    return {"status": "started", "scope": "all active companies"}


@app.post("/api/scrape/test")
def scrape_test():
    """Validation endpoint: scrape Anthropic only, synchronously.

    Runs inline (not threaded) so the caller gets the result directly — this is
    the 'test one before scaling' checkpoint from the build plan.
    """
    anthropic = database.get_company_by_name("Anthropic")
    if not anthropic:
        raise HTTPException(404, "Anthropic not found in companies table")
    result = get_orchestrator().run_full_scrape(
        only_company_id=anthropic["id"], note="test-anthropic"
    )
    return result


@app.post("/api/scrape/{company_id}")
def scrape_one(company_id: int):
    """Scrape a single company (async)."""
    comp = database.get_company(company_id)
    if not comp:
        raise HTTPException(404, "company not found")
    _trigger_async(only_company_id=company_id, note=f"single:{comp['name']}")
    return {"status": "started", "company": comp["name"]}


@app.get("/api/scrape/status")
def scrape_status():
    return SCRAPE_STATUS


# -----------------------------------------------------------------------------
# Job endpoints
# -----------------------------------------------------------------------------
@app.get("/api/jobs")
def get_jobs(match_only: bool = False, new_only: bool = False,
             tier: Optional[int] = None, city: Optional[str] = None, us_only: bool = True):
    return database.query_jobs(
        match_only=match_only, new_only=new_only, tier=tier,
        city=city, us_only=us_only,
    )


@app.get("/api/jobs/matches")
def get_matches():
    """Shortcut: matched + US-based only."""
    return database.query_jobs(match_only=True, us_only=True)


@app.get("/api/jobs/{job_id}/description")
def get_job_description(job_id: int):
    job = database.get_job(job_id)
    if not job:
        raise HTTPException(404, "job not found")
    return {"id": job_id, "title": job["title"], "description": job.get("description") or ""}


@app.post("/api/jobs/{job_id}/analyze")
def analyze_job(job_id: int, body: ResumeIn):
    """Run ResumeAnalysisAgent on a stored job's JD vs the supplied resume.

    If no resume_text is provided in the request body, falls back to the
    stored base resume (uploaded via Config > Resume Documents).
    """
    job = database.get_job(job_id)
    if not job:
        raise HTTPException(404, "job not found")
    orch = get_orchestrator()
    jd_text = orch.ensure_description(job_id) or job.get("title", "")

    resume_text = (body.resume_text or "").strip()
    if not resume_text:
        stored = database.get_document("base_resume")
        if stored:
            resume_text = stored["text_content"]
    if not resume_text:
        raise HTTPException(400, "No resume text provided and no stored resume found — upload one in Config > Resume Documents")

    analysis = orch.analyze_resume(jd_text, resume_text)
    return analysis


# -----------------------------------------------------------------------------
# Resume upload endpoint (per-job ad-hoc upload, kept for backwards compat)
# -----------------------------------------------------------------------------
@app.post("/api/resume/upload")
async def upload_resume(file: UploadFile = File(...)):
    """Extract plain text from an uploaded PDF or DOCX resume."""
    filename = file.filename or ""
    data = await file.read()

    if filename.lower().endswith(".pdf"):
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif filename.lower().endswith(".docx"):
        import io
        from docx import Document
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        raise HTTPException(400, "Only .pdf and .docx files are supported")

    return JSONResponse({"text": text.strip()})


# -----------------------------------------------------------------------------
# Master CV + base resume storage
# -----------------------------------------------------------------------------
def _extract_text(data: bytes, filename: str) -> str:
    """Extract plain text from PDF or DOCX bytes."""
    import io
    if filename.lower().endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    elif filename.lower().endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    raise HTTPException(400, "Only .pdf and .docx files are supported")


@app.post("/api/documents/cv")
async def upload_cv(file: UploadFile = File(...)):
    """Store master CV text (used as the content pool for resume generation)."""
    data = await file.read()
    text = _extract_text(data, file.filename or "")
    database.save_document("cv", text)
    return {"ok": True, "chars": len(text)}


@app.post("/api/documents/base-resume")
async def upload_base_resume(file: UploadFile = File(...)):
    """Store base resume as text + original DOCX bytes (used as the generation template)."""
    filename = file.filename or ""
    if not filename.lower().endswith(".docx"):
        raise HTTPException(400, "Base resume must be a .docx file so it can be used as a template")
    data = await file.read()
    text = _extract_text(data, filename)
    database.save_document("base_resume", text, binary_content=data)
    return {"ok": True, "chars": len(text)}


@app.get("/api/documents/status")
def documents_status():
    return database.documents_status()


@app.post("/api/jobs/{job_id}/generate")
async def generate_resume(job_id: int):
    """Generate a tailored resume DOCX for a specific job, using the stored CV + base resume."""
    from agents import resume_builder

    cv_doc = database.get_document("cv")
    base_doc = database.get_document("base_resume")
    if not cv_doc:
        raise HTTPException(400, "No master CV uploaded — go to Config > Resume Documents")
    if not base_doc or not base_doc.get("binary_content"):
        raise HTTPException(400, "No base resume uploaded — go to Config > Resume Documents")

    job = database.get_job(job_id)
    if not job:
        raise HTTPException(404, "job not found")

    orch = get_orchestrator()
    jd_text = orch.ensure_description(job_id) or job.get("title", "")

    content = resume_builder.generate_content(
        jd_text=jd_text,
        cv_text=cv_doc["text_content"],
        base_resume_text=base_doc["text_content"],
    )

    docx_bytes = resume_builder.build_docx(
        base_docx_bytes=base_doc["binary_content"],
        content_json=content,
    )

    safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in job.get("title", "resume"))
    filename = f"Resume_{safe_title[:40].strip()}.docx"

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# -----------------------------------------------------------------------------
# Run history
# -----------------------------------------------------------------------------
@app.get("/api/runs")
def get_runs():
    return database.list_runs()


# -----------------------------------------------------------------------------
# Frontend (served from /static, index at root)
# -----------------------------------------------------------------------------
app.mount("/static", StaticFiles(directory="static"), name="static")


@app.get("/")
def index():
    return FileResponse("static/index.html")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=False)
